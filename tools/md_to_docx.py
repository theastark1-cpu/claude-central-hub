#!/usr/bin/env python3
"""Convert the Monily cover memo markdown to a formatted DOCX.

Handles: H1/H2/H3, bullet lists, pipe tables, code blocks, bold (**),
italic (*), inline code (`), and horizontal rules.

Usage:
    python tools/md_to_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

REPO_ROOT = Path(__file__).resolve().parent.parent
SRC = REPO_ROOT / "monily-package" / "05_Cover_Memo_for_Monily.md"
DST = REPO_ROOT / "monily-package" / "05_Cover_Memo_for_Monily.docx"


def add_inline(paragraph, text: str, base_bold: bool = False, base_italic: bool = False):
    """Add text with inline **bold**, *italic*, `code` formatting."""
    # Tokenize: split on **bold**, *italic*, `code`, leaving plain text in between.
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if base_italic:
                run.italic = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            if base_bold:
                run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            if base_bold:
                run.bold = True
            if base_italic:
                run.italic = True


def shade_cell(cell, color_hex: str):
    """Apply background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def parse_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    """Parse a markdown pipe table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # Skip the separator line (---|---|---)
        if re.match(r'^\|[\s\-:|]+\|?\s*$', line):
            i += 1
            continue
        # Strip leading/trailing pipes and split
        line_inner = line.strip("|")
        cells = [c.strip() for c in line_inner.split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(src: Path, dst: Path):
    text = src.read_text()
    lines = text.split("\n")

    doc = Document()

    # Set default font + margins
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End: emit accumulated code
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                shade_cell_for_paragraph(p, "F2F2F2")
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)
            # Use a bottom border
            p_pr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            p_pr.append(pBdr)
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(level=0)
            run = h.add_run(stripped[2:])
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(level=1)
            run = h.add_run(stripped[3:])
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(level=2)
            run = h.add_run(stripped[4:])
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|?\s*$', lines[i+1].strip()):
            rows, end_idx = parse_table(lines, i)
            if rows:
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = 'Light Grid Accent 1'
                tbl.autofit = True
                for ridx, rcells in enumerate(rows):
                    for cidx, cell_text in enumerate(rcells):
                        cell = tbl.rows[ridx].cells[cidx]
                        cell.text = ""  # clear default
                        para = cell.paragraphs[0]
                        add_inline(para, cell_text, base_bold=(ridx == 0))
                        if ridx == 0:
                            shade_cell(cell, "1F4E79")
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.bold = True
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            doc.add_paragraph()  # spacing after table
            i = end_idx
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            content = re.sub(r'^\d+\.\s', '', stripped)
            add_inline(p, content)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Bold-only line (e.g., **Key:** value)
        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(dst)
    print(f"Wrote {dst}")


def shade_cell_for_paragraph(paragraph, color_hex: str):
    """Apply background shading to a paragraph (for code blocks)."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)


if __name__ == "__main__":
    md_to_docx(SRC, DST)
